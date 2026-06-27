#!/usr/bin/env python3
"""Generate downloadable resume files from one YAML source."""

from __future__ import annotations

import argparse
from calendar import monthrange
from datetime import date
from functools import lru_cache
import hashlib
import html
import json
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any
from urllib.parse import quote, urljoin, urlsplit

try:
    from babel.dates import format_date
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing date formatting dependency. Install Babel or run through scripts/build_resume_formats.sh."
    ) from exc

try:
    import yaml
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing YAML dependency. Install PyYAML or run through scripts/build_resume_formats.sh."
    ) from exc

try:
    import jsonschema
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing JSON Schema dependency. Install jsonschema or run through scripts/build_resume_formats.sh."
    ) from exc

try:
    from jinja2 import Environment, FileSystemLoader, select_autoescape
    from markupsafe import Markup
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing HTML templating dependency. Install Jinja2 or run through scripts/build_resume_formats.sh."
    ) from exc

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing DOCX dependencies. Install python-docx or run through scripts/build_resume_formats.sh."
    ) from exc

try:
    from weasyprint import HTML
except ImportError as exc:  # pragma: no cover - handled at runtime for local setup clarity.
    raise SystemExit(
        "Missing PDF dependencies. Install WeasyPrint or run through scripts/build_resume_formats.sh."
    ) from exc

ROOT = Path(__file__).resolve().parents[1]
MEDIA_ROOT = ROOT / "assets/media"
TEMPLATE_ROOT = ROOT / "scripts/templates"
PUBLIC_ASSET_ROOT = "../assets"
PUBLIC_MEDIA_ROOT = f"{PUBLIC_ASSET_ROOT}/media"
LANGUAGE_PATH_TEMPLATE = "/{lang}/"
DOWNLOAD_PATH_TEMPLATE = "/{lang}/{file}"
SITE_COVER_IMAGE = "/assets/og-cover-recruiter.jpg"
YANDEX_METRIKA_ID = "109268996"
YANDEX_METRIKA_ORIGIN = "https://mc.yandex.ru"
RESPONSIVE_IMAGE_DIR_NAME = "generated"
RESPONSIVE_IMAGE_WIDTHS = (480, 720, 1080)
RESPONSIVE_IMAGE_FORMATS = ("avif", "webp")
RESPONSIVE_SOURCE_EXTENSIONS = (".jpg", ".jpeg", ".png")
HERO_PRELOAD_WIDTH = RESPONSIVE_IMAGE_WIDTHS[1]
HERO_PRELOAD_FORMAT = RESPONSIVE_IMAGE_FORMATS[0]
THEMED_MEDIA_FILES = {
    "briefcase.svg",
    "contact.svg",
    "download.svg",
    "education.svg",
    "exnode.svg",
    "external-link.svg",
    "github.svg",
    "layers.svg",
    "location.svg",
    "moon.svg",
    "sbertech.svg",
    "star.svg",
    "sun.svg",
    "vstu.svg",
}

DOWNLOAD_STYLES = {
    "colors": {
        "accent": "1D5FD1",
        "text": "0F1F36",
        "muted": "526581",
        "tagBg": "E7EFFB",
        "tagText": "1649A3",
    },
    "typography": {
        "title": {"fontSize": 25.0, "lineHeight": 29.0},
        "headline": {"fontSize": 12.2, "lineHeight": 14.6},
        "contact": {"fontSize": 9.9, "lineHeight": 12.2},
        "summary": {"fontSize": 11.0, "lineHeight": 14.1},
        "section": {"fontSize": 12.6, "lineHeight": 14.6},
        "company": {"fontSize": 11.8, "lineHeight": 14.2},
        "meta": {"fontSize": 9.9, "lineHeight": 12.2},
        "body": {"fontSize": 11.0, "lineHeight": 14.1},
        "bullet": {"fontSize": 10.6, "lineHeight": 13.6},
    },
    "layout": {
        "pageMargin": 0.45,
        "maxPages": 2,
        "scaleSearchIterations": 12,
        "scaleSafetyMargin": 0.003,
        "minScale": {
            "en": 0.76,
            "ru": 0.72,
        },
        "maxScale": {
            "en": 1.16,
            "ru": 1.12,
        },
        "docxScaleAdjustment": {
            "en": 1.0,
            "ru": 1.0,
        },
    },
    "spacing": {
        "titleAfter": 2.0,
        "contactAfter": 2.6,
        "headlineAfter": 3.4,
        "sectionBefore": 5.4,
        "sectionAfter": 1.8,
        "entryBefore": 6.2,
        "entryHeaderAfter": 1.6,
        "paragraphAfter": 2.6,
        "bulletAfter": 1.6,
        "tagGap": 4.0,
        "skillAfter": 1.0,
    },
}
DOWNLOAD_DOCX_FONT_NAME = "Arial"
DOWNLOAD_PDF_FONT_FAMILY = "ResumeDownloadSans"

DOWNLOAD_SECTION_SOURCE_KEYS = {
    "profile": "strengths",
    "experience": "experience",
    "education": "education",
    "skills": "skills",
}

def localized_tree(value: Any, lang: str, languages: list[str]) -> Any:
    if isinstance(value, dict) and set(value.keys()).issubset(set(languages)):
        return value[lang]
    if isinstance(value, dict):
        return {key: localized_tree(item, lang, languages) for key, item in value.items()}
    if isinstance(value, list):
        return [localized_tree(item, lang, languages) for item in value]
    return value


def load_yaml_mapping(path: Path) -> dict[str, Any]:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def html_attr(value: Any) -> str:
    return html.escape(str(value), quote=True)


def html_text(value: Any) -> str:
    return html.escape(str(value))


def slug_file_stem(file_name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", Path(file_name).stem.lower()).strip("-")


def public_media_path(file_name: str) -> str:
    return f"{PUBLIC_MEDIA_ROOT}/{file_name}"


def site_runtime_config() -> dict[str, Any]:
    return {
        "media": {
            "responsive": {
                "directory": RESPONSIVE_IMAGE_DIR_NAME,
                "widths": list(RESPONSIVE_IMAGE_WIDTHS),
                "formats": list(RESPONSIVE_IMAGE_FORMATS),
                "sourceExtensions": [extension.removeprefix(".") for extension in RESPONSIVE_SOURCE_EXTENSIONS],
            },
            "themedFiles": sorted(THEMED_MEDIA_FILES),
        },
    }


def analytics_runtime_config() -> dict[str, str]:
    return {
        "yandexMetrikaId": YANDEX_METRIKA_ID,
        "yandexMetrikaOrigin": YANDEX_METRIKA_ORIGIN,
    }


def runtime_page_data(source: dict[str, Any], lang: str) -> dict[str, Any]:
    return {
        "lang": lang,
        "languages": source["languages"],
        "site": {
            **source["site"],
            **site_runtime_config(),
        },
        "analytics": analytics_runtime_config(),
        "theme": localized_tree(source["siteUi"]["theme"], lang, source["languages"]),
    }


def responsive_variant_path(file_name: str, width: int, extension: str) -> str:
    return f"{PUBLIC_MEDIA_ROOT}/{RESPONSIVE_IMAGE_DIR_NAME}/{slug_file_stem(file_name)}-{width}.{extension}"


def responsive_srcset(file_name: str, extension: str) -> str:
    return ", ".join(
        f"{responsive_variant_path(file_name, width, extension)} {width}w"
        for width in RESPONSIVE_IMAGE_WIDTHS
    )


def relative_public_path(path: str, depth: int = 1) -> str:
    prefix = "../" * depth
    return f"{prefix}{path.lstrip('/')}"


def supports_responsive_variants(file_name: str | None) -> bool:
    return bool(file_name) and Path(str(file_name)).suffix.lower() in RESPONSIVE_SOURCE_EXTENSIONS


def responsive_picture_html(
    file_name: str,
    *,
    css_class: str,
    alt: str,
    sizes: str,
    loading: str = "lazy",
    fetchpriority: str | None = None,
    image_id: str | None = None,
    style: str | None = None,
) -> str:
    image_attributes = [
        f'class="{html_attr(css_class)}"',
        f'src="{html_attr(public_media_path(file_name))}"',
        f'alt="{html_attr(alt)}"',
        f'sizes="{html_attr(sizes)}"',
        'decoding="async"',
    ]
    if image_id:
        image_attributes.insert(0, f'id="{html_attr(image_id)}"')
    if loading:
        image_attributes.append(f'loading="{html_attr(loading)}"')
    if fetchpriority:
        image_attributes.append(f'fetchpriority="{html_attr(fetchpriority)}"')
    if style:
        image_attributes.append(f'style="{html_attr(style)}"')

    if supports_responsive_variants(file_name):
        sources = "\n".join(
            f'                <source type="image/{extension}" '
            f'srcset="{html_attr(responsive_srcset(file_name, extension))}" '
            f'sizes="{html_attr(sizes)}" />'
            for extension in RESPONSIVE_IMAGE_FORMATS
        )
        return (
            "              <picture>\n"
            f"{sources}\n"
            f"                <img {' '.join(image_attributes)} />\n"
            "              </picture>"
        )

    return f"              <img {' '.join(image_attributes)} />"


def absolute_site_url(source: dict[str, Any], path: str = "/") -> str:
    base = str(source["site"]["url"]).rstrip("/") + "/"
    return urljoin(base, path.lstrip("/"))


def format_site_path(template: str, **values: str) -> str:
    encoded_values = {key: quote(str(value), safe="") for key, value in values.items()}
    path = template.format(**encoded_values)
    return path if path.startswith("/") else f"/{path}"


def page_path(source: dict[str, Any], lang: str) -> str:
    return format_site_path(LANGUAGE_PATH_TEMPLATE, lang=lang)


def download_path(source: dict[str, Any], lang: str, file_name: str) -> str:
    return format_site_path(DOWNLOAD_PATH_TEMPLATE, lang=lang, file=file_name)


def page_url(source: dict[str, Any], lang: str) -> str:
    return absolute_site_url(source, page_path(source, lang))


def site_host(source: dict[str, Any]) -> str:
    return urlsplit(str(source["site"]["url"])).netloc


def collect_photo_file_names(source: dict[str, Any]) -> list[str]:
    files: list[str] = []
    photo = source.get("person", {}).get("photo", {})
    if photo.get("src"):
        files.append(photo["src"])
    return sorted(set(files))


def image_magick_command() -> list[str] | None:
    magick = shutil.which("magick")
    if magick:
        return [magick]
    convert = shutil.which("convert")
    if convert:
        return [convert]
    return None


def generate_responsive_media(source: dict[str, Any]) -> None:
    command = image_magick_command()
    photo_files = [
        file_name
        for file_name in collect_photo_file_names(source)
        if supports_responsive_variants(file_name)
    ]
    if not photo_files:
        return
    if not command:
        raise RuntimeError("ImageMagick is required to generate AVIF/WebP responsive photo variants.")

    output_dir = MEDIA_ROOT / RESPONSIVE_IMAGE_DIR_NAME
    output_dir.mkdir(parents=True, exist_ok=True)
    expected_outputs: set[Path] = set()
    for file_name in photo_files:
        source_path = MEDIA_ROOT / file_name
        if not source_path.exists():
            raise RuntimeError(f"Referenced photo asset does not exist: {source_path.relative_to(ROOT)}")
        for width in RESPONSIVE_IMAGE_WIDTHS:
            for extension in RESPONSIVE_IMAGE_FORMATS:
                output_path = output_dir / f"{slug_file_stem(file_name)}-{width}.{extension}"
                expected_outputs.add(output_path)
                if output_path.exists() and output_path.stat().st_mtime >= source_path.stat().st_mtime:
                    continue
                quality = "52" if extension == "avif" else "82"
                subprocess.run(
                    [
                        *command,
                        str(source_path),
                        "-auto-orient",
                        "-resize",
                        f"{width}x>",
                        "-strip",
                        "-quality",
                        quality,
                        str(output_path),
                    ],
                    check=True,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.PIPE,
                    text=True,
                )
    for existing in output_dir.iterdir():
        if existing.is_file() and existing.suffix.lower().removeprefix(".") in RESPONSIVE_IMAGE_FORMATS:
            if existing not in expected_outputs:
                existing.unlink()


def hex_color(value: str) -> str:
    return str(value).strip().removeprefix("#").upper()


def rgb_color(hex_value: str) -> RGBColor:
    color = hex_color(hex_value)
    return RGBColor(
        int(color[0:2], 16),
        int(color[2:4], 16),
        int(color[4:6], 16),
    )


def parse_resume_date(value: str, *, upper_bound: bool) -> date:
    parts = value.split("-")
    year = int(parts[0])
    if len(parts) == 1:
        month = 12 if upper_bound else 1
        day = 31 if upper_bound else 1
        return date(year, month, day)
    month = int(parts[1])
    day = monthrange(year, month)[1] if upper_bound else 1
    return date(year, month, day)


def months_between(start_date: date, end_date: date) -> int:
    months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
    if end_date.day < start_date.day:
        months -= 1
    return max(months, 0)


def format_duration_label(total_months: int, lang: str) -> str:
    years, months = divmod(total_months, 12)
    parts: list[str] = []
    if years:
        if lang == "ru":
            parts.append(f"{years} г.")
        else:
            parts.append(f"{years} {'yr' if years == 1 else 'yrs'}")
    if months or not parts:
        if lang == "ru":
            parts.append(f"{months} мес.")
        else:
            parts.append(f"{months} {'mo' if months == 1 else 'mos'}")
    return " ".join(parts)


def experience_duration(item: dict[str, Any], lang: str) -> str:
    end_date = parse_resume_date(item["endDate"], upper_bound=True) if item.get("endDate") else date.today()
    start_date = parse_resume_date(item["startDate"], upper_bound=False)
    return format_duration_label(months_between(start_date, end_date), lang)


def contact_value(contact: dict[str, str]) -> str:
    return str(contact["value"]).strip()


def is_web_url(value: str) -> bool:
    parsed = urlsplit(value)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def is_email_address(value: str) -> bool:
    local, separator, domain = value.partition("@")
    return (
        bool(local and separator and domain)
        and "." in domain
        and not any(character.isspace() for character in value)
    )


def contact_href(contact: dict[str, str]) -> str | None:
    value = contact_value(contact)
    if is_web_url(value):
        return value
    if is_email_address(value):
        return f"mailto:{value}"
    return None


def contact_link_text(contact: dict[str, str], *, shorten_web_urls: bool = True) -> str:
    value = contact_value(contact)
    if shorten_web_urls and is_web_url(value):
        return value.removeprefix("https://").removeprefix("http://")
    return value


def is_safe_media_file_name(file_name: str) -> bool:
    path = Path(file_name)
    return not path.is_absolute() and path.name == file_name and ".." not in path.parts


def media_source(file_name: str) -> Path:
    if not is_safe_media_file_name(file_name):
        return MEDIA_ROOT / "__invalid_media_reference__"
    for path in (MEDIA_ROOT / file_name, MEDIA_ROOT / "light" / file_name, MEDIA_ROOT / "dark" / file_name):
        if path.exists():
            return path
    return MEDIA_ROOT / file_name


def download_section_title(source: dict[str, Any], lang: str, key: str) -> str:
    section_key = DOWNLOAD_SECTION_SOURCE_KEYS[key]
    return localized_tree(source[section_key]["title"], lang, source["languages"])


def skill_group_segments(groups: list[dict[str, Any]]) -> list[tuple[str, str]]:
    return [(group["title"], ", ".join(group["items"])) for group in groups]


def format_resume_date(value: str, lang: str) -> str:
    parts = value.split("-")
    if len(parts) == 1:
        return parts[0]
    pattern = "LLL y" if lang == "ru" else "MMM y"
    return format_date(parse_resume_date(value, upper_bound=False), pattern, locale=lang)


def format_period(item: dict[str, Any], labels: dict[str, str], lang: str) -> str:
    start = format_resume_date(item["startDate"], lang)
    end_date = item.get("endDate")
    end = format_resume_date(end_date, lang) if end_date else labels["present"]
    return f"{start} - {end}"


def resume_date_upper_bound(value: str | None) -> date | None:
    if value is None:
        return None
    return parse_resume_date(value, upper_bound=True)


def is_expected_education(item: dict[str, Any]) -> bool:
    return resume_date_upper_bound(item["endDate"]) > date.today()


def format_education_period(item: dict[str, Any], labels: dict[str, str], lang: str) -> str:
    start = format_resume_date(item["startDate"], lang)
    end_date = format_resume_date(item["endDate"], lang)
    end = f"{labels['expectedGraduation']}: {end_date}" if is_expected_education(item) else end_date
    return f"{start} - {end}"


def experience_header_text(item: dict[str, Any], labels: dict[str, str], lang: str) -> str:
    period = format_period(item, labels, lang)
    duration = experience_duration(item, lang)
    return f"{item['role']} ({period}, {duration})"


def experience_tag_values(item: dict[str, Any]) -> list[str]:
    return [value for value in (item.get("location"), item.get("employmentType")) if value]


def education_header_text(item: dict[str, Any], labels: dict[str, str], lang: str) -> str:
    return f"{item['degree']} ({format_education_period(item, labels, lang)})"


def download_file_names(source: dict[str, Any]) -> dict[str, str]:
    file_base_name = resume_file_base_name(source)
    return {
        "pdf": f"{file_base_name}.pdf",
        "docx": f"{file_base_name}.docx",
        "txt": f"{file_base_name}.txt",
    }


def resume_file_base_name(source: dict[str, Any]) -> str:
    lang = source.get("defaultLanguage", source["languages"][0])
    localized_name = localized_tree(source["person"]["name"], lang, source["languages"])
    return "_".join(localized_name.split())


def schema_error_path(error: Any) -> str:
    path = "/".join(str(part) for part in error.absolute_path)
    return f"/{path}" if path else "/"


def validate_json_schema(source: dict[str, Any], schema: dict[str, Any]) -> None:
    validator = jsonschema.Draft7Validator(schema)
    errors = sorted(validator.iter_errors(source), key=lambda item: list(item.absolute_path))
    if errors:
        error = errors[0]
        raise ValueError(f"Schema validation failed at {schema_error_path(error)}: {error.message}")


def validate_date_order(items: list[dict[str, Any]], path: str) -> None:
    for index, item in enumerate(items):
        end_date = item.get("endDate")
        if not end_date:
            continue
        if parse_resume_date(end_date, upper_bound=True) < parse_resume_date(item["startDate"], upper_bound=False):
            raise ValueError(f"Schema custom validation failed at {path}/{index}: endDate must be >= startDate.")


def validate_experience_sorting(source: dict[str, Any]) -> None:
    starts = [item["startDate"] for item in source["experience"]["items"]]
    if starts != sorted(starts, reverse=True):
        raise ValueError("Schema custom validation failed at /experience/items: items must be sorted by startDate descending.")


def walk_values(value: Any) -> Any:
    yield value
    if isinstance(value, dict):
        for item in value.values():
            yield from walk_values(item)
    elif isinstance(value, list):
        for item in value:
            yield from walk_values(item)


def validate_localized_values(source: dict[str, Any]) -> None:
    languages = set(source["languages"])
    for value in walk_values(source):
        if isinstance(value, dict) and languages.intersection(value.keys()):
            if set(value.keys()) != languages:
                raise ValueError("Schema custom validation failed: localized maps must cover all configured languages.")


def collect_asset_paths(value: Any) -> list[str]:
    paths: list[str] = []
    if isinstance(value, dict):
        for key, item in value.items():
            if key in {"icon", "src"} and isinstance(item, str):
                paths.append(item)
            else:
                paths.extend(collect_asset_paths(item))
    elif isinstance(value, list):
        for item in value:
            paths.extend(collect_asset_paths(item))
    return paths


def validate_asset_paths(source: dict[str, Any]) -> None:
    for asset in sorted(set(collect_asset_paths(source))):
        if not media_source(asset).exists():
            raise ValueError(f"Schema custom validation failed: asset does not exist under assets/media: {asset}")


def validate_custom_schema_rules(source: dict[str, Any], schema: dict[str, Any]) -> None:
    rules = {item.get("rule") for item in schema.get("x-jorqenCustomValidations", [])}
    if "sortedByStartDateDescending" in rules:
        validate_experience_sorting(source)
    if "endDateGreaterThanOrEqualStartDate" in rules:
        validate_date_order(source["experience"]["items"], "/experience/items")
        validate_date_order(source["education"]["items"], "/education/items")
    if "localizedValuesMustCoverConfiguredLanguages" in rules:
        validate_localized_values(source)
    if "assetPathMustExistUnderAssetsMedia" in rules:
        validate_asset_paths(source)


def validate_source(data_path: Path, source: dict[str, Any]) -> None:
    schema_path = data_path.with_name("resume.schema.yaml")
    if not schema_path.exists():
        raise ValueError(f"Missing schema file: {schema_path}")
    schema = load_yaml_mapping(schema_path)
    validate_json_schema(source, schema)
    validate_custom_schema_rules(source, schema)


def load_data(path: Path) -> dict[str, Any]:
    source = load_yaml_mapping(path)
    validate_source(path, source)
    source["contacts"]["website"]["value"] = source["site"]["url"].rstrip("/")
    return source


def contact_keys(source: dict[str, Any]) -> list[str]:
    return list(source["contacts"].keys())


def contact(source: dict[str, Any], key: str, lang: str) -> dict[str, str]:
    return localized_tree(source["contacts"][key], lang, source["languages"])


def person(source: dict[str, Any], lang: str) -> dict[str, Any]:
    return localized_tree(source["person"], lang, source["languages"])


def labels(source: dict[str, Any], lang: str) -> dict[str, str]:
    return localized_tree(source["resumeLabels"], lang, source["languages"])


def section(source: dict[str, Any], key: str, lang: str) -> dict[str, Any]:
    return localized_tree(source[key], lang, source["languages"])


def experience_items(source: dict[str, Any], lang: str) -> list[dict[str, Any]]:
    return sorted(
        section(source, "experience", lang)["items"],
        key=lambda item: resume_date_upper_bound(item["startDate"]),
        reverse=True,
    )


def contact_parts(source: dict[str, Any], lang: str) -> list[str]:
    return [
        contact_link_text(contact(source, key, lang), shorten_web_urls=False)
        for key in contact_keys(source)
    ]


def add_txt_line(lines: list[str], text: str = "") -> None:
    if text:
        lines.extend(str(text).splitlines())
    else:
        lines.append("")


def generate_txt(source: dict[str, Any], lang: str, output_path: Path) -> None:
    profile = person(source, lang)
    label_values = labels(source, lang)
    experience = section(source, "experience", lang)
    education = section(source, "education", lang)
    skills = section(source, "skills", lang)
    lines: list[str] = []

    lines.append(profile["name"])
    lines.append("=" * len(profile["name"]))
    lines.append(" | ".join(contact_parts(source, lang)))
    add_txt_line(lines)
    add_txt_line(lines, profile["role"])

    add_txt_line(lines)
    lines.append(download_section_title(source, lang, "profile"))
    lines.append("-" * len(download_section_title(source, lang, "profile")))
    add_txt_line(lines, profile["summary"])

    add_txt_line(lines)
    lines.append(download_section_title(source, lang, "experience"))
    lines.append("-" * len(download_section_title(source, lang, "experience")))
    for item in experience_items(source, lang):
        add_txt_line(lines)
        header_line = f"{item['company']} | {experience_header_text(item, label_values, lang)}"
        tags = experience_tag_values(item)
        if tags:
            header_line = f"{header_line} · {' · '.join(tags)}"
        lines.append(header_line)
        if item.get("url"):
            lines.append(f"{experience['companySiteLabel']}: {item['url']}")
        add_txt_line(lines, item["summary"])
        for bullet in item["highlights"]:
            for index, bullet_line in enumerate(str(bullet).splitlines()):
                prefix = "- " if index == 0 else "  "
                lines.append(f"{prefix}{bullet_line}")

    add_txt_line(lines)
    lines.append(download_section_title(source, lang, "education"))
    lines.append("-" * len(download_section_title(source, lang, "education")))
    for item in education["items"]:
        add_txt_line(lines)
        lines.append(f"{item['institution']} | {education_header_text(item, label_values, lang)}")
        education_site_label = education.get("institutionSiteLabel")
        if education_site_label and item.get("url"):
            lines.append(f"{education_site_label}: {item['url']}")

    add_txt_line(lines)
    lines.append(download_section_title(source, lang, "skills"))
    lines.append("-" * len(download_section_title(source, lang, "skills")))
    for title, values in skill_group_segments(skills["groups"]):
        lines.append(f"{title}: {values}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def add_hyperlink(paragraph: Any, text: str, url: str, color_hex: str, bold: bool = False) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Arial")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    properties.append(fonts)
    properties.append(color)
    properties.append(underline)
    if bold:
        properties.append(OxmlElement("w:b"))
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # noqa: SLF001 - python-docx exposes no public hyperlink writer.


def add_docx_tag_run(
    paragraph: Any,
    text: str,
    *,
    font_name: str,
    font_size: float,
    fill_hex: str,
    text_hex: str,
) -> None:
    run = paragraph.add_run(f" {text} ")
    run.bold = False
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = rgb_color(text_hex)
    properties = run._element.get_or_add_rPr()  # noqa: SLF001
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)
    properties.append(fonts)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color(fill_hex))
    properties.append(shading)


def docx_paragraph_style(styles: Any, name: str) -> Any:
    try:
        return styles[name]
    except KeyError:
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_docx_style_font_family(style: Any, font_name: str) -> None:
    properties = style._element.get_or_add_rPr()  # noqa: SLF001
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.append(fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)


def set_docx_style_widow_control(style: Any, enabled: bool = False) -> None:
    paragraph_properties = style._element.get_or_add_pPr()  # noqa: SLF001
    widow_control = paragraph_properties.find(qn("w:widowControl"))
    if widow_control is None:
        widow_control = OxmlElement("w:widowControl")
        paragraph_properties.append(widow_control)
    widow_control.set(qn("w:val"), "1" if enabled else "0")


def configure_docx_paragraph_style(
    styles: Any,
    name: str,
    font_name: str,
    font_size: float,
    color: RGBColor,
    *,
    bold: bool = False,
    alignment: Any = None,
    keep_with_next: bool = False,
    keep_together: bool = False,
    line_spacing: float = 12,
    space_before: float = 0,
    space_after: float = 0,
    left_indent: float | None = None,
    first_line_indent: float | None = None,
) -> None:
    style = docx_paragraph_style(styles, name)
    style.font.name = font_name
    set_docx_style_font_family(style, font_name)
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.font.color.rgb = color
    style.paragraph_format.alignment = alignment
    style.paragraph_format.keep_with_next = keep_with_next
    style.paragraph_format.keep_together = keep_together
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.line_spacing = Pt(line_spacing)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    set_docx_style_widow_control(style, enabled=False)
    if left_indent is not None:
        style.paragraph_format.left_indent = Inches(left_indent)
    if first_line_indent is not None:
        style.paragraph_format.first_line_indent = Inches(first_line_indent)


def scaled_typography(download_style: dict[str, Any], scale: float) -> dict[str, dict[str, float]]:
    return {
        key: {
            "fontSize": values["fontSize"] * scale,
            "lineHeight": values["lineHeight"] * scale,
        }
        for key, values in download_style["typography"].items()
    }


def configure_docx(document: Document, download_style: dict[str, Any], scale: float) -> None:
    style_colors = download_style["colors"]
    typography = scaled_typography(download_style, scale)
    layout = download_style["layout"]
    spacing = download_style["spacing"]
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    page_margin = Inches(layout["pageMargin"])
    section.top_margin = page_margin
    section.bottom_margin = page_margin
    section.left_margin = page_margin
    section.right_margin = page_margin
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    font_name = DOWNLOAD_DOCX_FONT_NAME
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    set_docx_style_font_family(normal, font_name)
    normal.font.size = Pt(typography["body"]["fontSize"])
    normal.font.color.rgb = rgb_color(style_colors["text"])
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(typography["body"]["lineHeight"])
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    set_docx_style_widow_control(normal, enabled=False)

    configure_docx_paragraph_style(
        styles,
        "ResumeTitle",
        font_name,
        typography["title"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        line_spacing=typography["title"]["lineHeight"],
        space_after=spacing["titleAfter"],
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeHeadline",
        font_name,
        typography["headline"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        line_spacing=typography["headline"]["lineHeight"],
        space_after=spacing["headlineAfter"],
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeContact",
        font_name,
        typography["contact"]["fontSize"],
        rgb_color(style_colors["muted"]),
        line_spacing=typography["contact"]["lineHeight"],
        space_after=spacing["contactAfter"],
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeSummary",
        font_name,
        typography["summary"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["summary"]["lineHeight"],
        space_after=0,
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeSection",
        font_name,
        typography["section"]["fontSize"],
        rgb_color(style_colors["accent"]),
        bold=True,
        keep_with_next=True,
        line_spacing=typography["section"]["lineHeight"],
        space_before=spacing["sectionBefore"],
        space_after=spacing["sectionAfter"],
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeCompany",
        font_name,
        typography["company"]["fontSize"],
        rgb_color(style_colors["text"]),
        bold=True,
        keep_with_next=True,
        line_spacing=typography["company"]["lineHeight"],
        space_before=spacing["entryBefore"],
        space_after=spacing["entryHeaderAfter"],
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeBody",
        font_name,
        typography["body"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["body"]["lineHeight"],
        space_after=spacing["paragraphAfter"],
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeSkill",
        font_name,
        typography["body"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["body"]["lineHeight"],
        space_after=spacing["skillAfter"],
    )
    configure_docx_paragraph_style(
        styles,
        "ResumeBullet",
        font_name,
        typography["bullet"]["fontSize"],
        rgb_color(style_colors["text"]),
        line_spacing=typography["bullet"]["lineHeight"],
        space_after=spacing["bulletAfter"],
        left_indent=0.17,
        first_line_indent=-0.11,
    )


def docx_contact_line(
    paragraph: Any,
    source: dict[str, Any],
    lang: str,
    accent_color: str,
) -> None:
    for index, key in enumerate(contact_keys(source)):
        contact_value_for_lang = contact(source, key, lang)
        if index:
            paragraph.add_run("  |  ")
        href = contact_href(contact_value_for_lang)
        if href:
            add_hyperlink(paragraph, contact_link_text(contact_value_for_lang), href, accent_color)
        else:
            paragraph.add_run(contact_link_text(contact_value_for_lang))


def generate_docx(source: dict[str, Any], lang: str, output_path: Path, build_dir: Path, scale: float) -> None:
    profile = person(source, lang)
    label_values = labels(source, lang)
    education = section(source, "education", lang)
    skills = section(source, "skills", lang)
    download_style = DOWNLOAD_STYLES
    style_colors = download_style["colors"]
    layout = download_style["layout"]
    typography = scaled_typography(download_style, scale)
    document = Document()
    configure_docx(document, download_style, scale)
    document.core_properties.title = f"{profile['name']} | {profile['headline']}"
    document.core_properties.author = profile["name"]
    document.core_properties.comments = "generated by scripts/generate_resume_outputs.py"

    document.add_paragraph(profile["name"], style="ResumeTitle")
    contact_paragraph = document.add_paragraph(style="ResumeContact")
    docx_contact_line(
        contact_paragraph,
        source,
        lang,
        style_colors["accent"],
    )
    document.add_paragraph(profile["role"], style="ResumeHeadline")

    document.add_paragraph(download_section_title(source, lang, "profile").upper(), style="ResumeSection")
    document.add_paragraph(profile["summary"], style="ResumeBody")

    document.add_paragraph(download_section_title(source, lang, "experience").upper(), style="ResumeSection")
    for item in experience_items(source, lang):
        header = document.add_paragraph(style="ResumeCompany")
        if item.get("url"):
            add_hyperlink(header, item["company"], item["url"], style_colors["accent"], bold=True)
        else:
            company_run = header.add_run(item["company"])
            company_run.bold = True
        header.add_run(f" | {experience_header_text(item, label_values, lang)}").bold = True
        for tag in experience_tag_values(item):
            header.add_run(" ").bold = False
            add_docx_tag_run(
                header,
                tag,
                font_name=DOWNLOAD_DOCX_FONT_NAME,
                font_size=typography["meta"]["fontSize"],
                fill_hex=style_colors["tagBg"],
                text_hex=style_colors["tagText"],
            )

        document.add_paragraph(item["summary"], style="ResumeBody")
        for bullet in item["highlights"]:
            document.add_paragraph(f"• {bullet}", style="ResumeBullet")

    document.add_paragraph(download_section_title(source, lang, "education").upper(), style="ResumeSection")
    for item in education["items"]:
        paragraph = document.add_paragraph(style="ResumeBody")
        if item.get("url"):
            add_hyperlink(paragraph, item["institution"], item["url"], style_colors["accent"], bold=True)
        else:
            paragraph.add_run(item["institution"]).bold = True
        paragraph.add_run(f" | {education_header_text(item, label_values, lang)}")

    document.add_paragraph(download_section_title(source, lang, "skills").upper(), style="ResumeSection")
    for title, values in skill_group_segments(skills["groups"]):
        skills_paragraph = document.add_paragraph(style="ResumeSkill")
        skills_paragraph.add_run(f"{title}: ").bold = True
        skills_paragraph.add_run(values)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def find_font(candidates: list[str]) -> Path | None:
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return path
    fc_match = shutil.which("fc-match")
    if fc_match:
        for family in ("Arial", "Liberation Sans", "DejaVu Sans", "Arial Unicode MS"):
            result = subprocess.run(
                [fc_match, "-f", "%{file}", family],
                check=False,
                text=True,
                capture_output=True,
            )
            if result.returncode == 0 and result.stdout.strip():
                path = Path(result.stdout.strip())
                if path.exists():
                    return path
    return None


def env_font_path(name: str) -> Path | None:
    value = os.environ.get(name)
    if not value:
        return None
    path = Path(value).expanduser()
    if not path.exists():
        raise RuntimeError(f"{name} points to a missing font file: {path}")
    return path


def download_font_paths() -> dict[str, str | None]:
    regular = env_font_path("RESUME_DOWNLOAD_FONT_REGULAR") or find_font([
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ])
    bold = env_font_path("RESUME_DOWNLOAD_FONT_BOLD") or find_font([
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial Bold.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]) or regular
    return {
        "regular": regular.resolve().as_uri() if regular else None,
        "bold": bold.resolve().as_uri() if bold else None,
    }


def download_contacts(source: dict[str, Any], lang: str, build_dir: Path) -> list[dict[str, Any]]:
    items = []
    for key in contact_keys(source):
        contact_value_for_lang = contact(source, key, lang)
        items.append({
            "key": key,
            "text": contact_link_text(contact_value_for_lang),
            "href": contact_href(contact_value_for_lang),
        })
    return items


def download_view_model(source: dict[str, Any], lang: str, build_dir: Path) -> dict[str, Any]:
    profile = person(source, lang)
    label_values = labels(source, lang)
    education = section(source, "education", lang)
    skills = section(source, "skills", lang)
    return {
        "title": f"{profile['name']} | {profile['headline']}",
        "profile": profile,
        "contacts": download_contacts(source, lang, build_dir),
        "sections": {
            "profile": {
                "title": download_section_title(source, lang, "profile"),
                "summary": profile["summary"],
            },
            "experience": {
                "title": download_section_title(source, lang, "experience"),
                "items": [
                    {
                        "company": item["company"],
                        "url": item.get("url"),
                        "header": experience_header_text(item, label_values, lang),
                        "tags": experience_tag_values(item),
                        "summary": item["summary"],
                        "highlights": item["highlights"],
                    }
                    for item in experience_items(source, lang)
                ],
            },
            "education": {
                "title": download_section_title(source, lang, "education"),
                "items": [
                    {
                        "institution": item["institution"],
                        "url": item.get("url"),
                        "header": education_header_text(item, label_values, lang),
                    }
                    for item in education["items"]
                ],
            },
            "skills": {
                "title": download_section_title(source, lang, "skills"),
                "groups": [
                    {"title": title, "values": values}
                    for title, values in skill_group_segments(skills["groups"])
                ],
            },
        },
    }


def render_download_html(source: dict[str, Any], lang: str, build_dir: Path, scale: float) -> str:
    return render_template(
        "resume_download.html.j2",
        lang=lang,
        model=download_view_model(source, lang, build_dir),
        colors=DOWNLOAD_STYLES["colors"],
        typography=scaled_typography(DOWNLOAD_STYLES, scale),
        layout=DOWNLOAD_STYLES["layout"],
        spacing=DOWNLOAD_STYLES["spacing"],
        font_family=DOWNLOAD_PDF_FONT_FAMILY,
        font_paths=download_font_paths(),
    )


def render_download_pdf_document(source: dict[str, Any], lang: str, build_dir: Path, scale: float) -> Any:
    return HTML(
        string=render_download_html(source, lang, build_dir, scale),
        base_url=str(ROOT),
    ).render()


def resolve_download_pdf(source: dict[str, Any], lang: str, build_dir: Path) -> tuple[float, Any]:
    layout = DOWNLOAD_STYLES["layout"]
    min_scale = layout["minScale"].get(lang, 0.78)
    max_scale = layout["maxScale"].get(lang, 1.2)
    max_pages = layout["maxPages"]

    min_document = render_download_pdf_document(source, lang, build_dir, min_scale)
    if len(min_document.pages) > max_pages:
        raise RuntimeError(
            f"Generated {lang.upper()} PDF has {len(min_document.pages)} pages at minimum scale {min_scale:.2f}; "
            f"target is {max_pages} pages."
        )

    max_document = render_download_pdf_document(source, lang, build_dir, max_scale)
    if len(max_document.pages) <= max_pages:
        return max_scale, max_document

    lower = min_scale
    upper = max_scale
    best_scale = min_scale
    for _ in range(layout["scaleSearchIterations"]):
        scale = round((lower + upper) / 2, 4)
        document = render_download_pdf_document(source, lang, build_dir, scale)
        if len(document.pages) <= max_pages:
            lower = scale
            best_scale = scale
        else:
            upper = scale

    scale = max(min_scale, round(best_scale - layout["scaleSafetyMargin"], 4))
    document = render_download_pdf_document(source, lang, build_dir, scale)
    while len(document.pages) > max_pages and scale > min_scale:
        scale = max(min_scale, round(scale - layout["scaleSafetyMargin"], 4))
        document = render_download_pdf_document(source, lang, build_dir, scale)
    return scale, document


def generate_pdf(source: dict[str, Any], lang: str, output_path: Path, build_dir: Path) -> float:
    scale, document = resolve_download_pdf(source, lang, build_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.write_pdf(str(output_path))
    return scale


def link_target_attributes(href: str | None) -> str:
    if not href:
        return ""
    if href.startswith(("mailto:", "tel:")):
        return ""
    return ' target="_blank" rel="noopener noreferrer"'


def render_icon(
    file_name: str,
    css_class: str = "",
    *,
    theme: str = "light",
    alt: str = "",
    hidden: bool = True,
) -> str:
    class_attr = f' class="{html_attr(css_class)}"' if css_class else ""
    aria = ' aria-hidden="true"' if hidden else ""
    data_media = f' data-media="{html_attr(file_name)}"' if file_name in THEMED_MEDIA_FILES else ""
    themed_prefix = f"{theme}/" if data_media else ""
    return (
        f'<img src="{html_attr(public_media_path(themed_prefix + file_name))}"{data_media}'
        f'{class_attr} alt="{html_attr(alt)}"{aria} />'
    )


def render_panel_icon(file_name: str) -> str:
    return f'<span class="panel-icon">{render_icon(file_name)}</span>'


def render_language_switch(source: dict[str, Any], lang: str) -> str:
    links = []
    for item_lang in source["languages"]:
        active = ' class="active"' if item_lang == lang else ""
        pressed = "true" if item_lang == lang else "false"
        path = relative_public_path(page_path(source, item_lang))
        links.append(
            f'<a href="{html_attr(path)}" data-lang-switch="{html_attr(item_lang)}" '
            f'data-lang-switch-url="{html_attr(path)}" aria-pressed="{pressed}"{active}>'
            f"{html_text(item_lang.upper())}</a>"
        )
    return "\n          ".join(links)


def render_contact_link(key: str, contact_data: dict[str, str], theme: str = "light") -> str:
    href = contact_href(contact_data) or ""
    label = contact_link_text(contact_data)
    value = contact_value(contact_data)
    target = link_target_attributes(href)
    return (
        f'<a id="hero-{html_attr(key)}" class="action-link" href="{html_attr(href)}"{target} '
        f'title="{html_attr(value)}" data-analytics-contact="true" '
        f'data-analytics-label="{html_attr(key.title())}" data-analytics-section="contacts">'
        f'{render_icon(contact_data.get("icon", ""), theme=theme) if contact_data.get("icon") else ""}'
        f"<span>{html_text(label)}</span></a>"
    )


def render_facts(items: list[dict[str, Any]]) -> str:
    cards = []
    for item in items:
        lines = [
            '<article class="fact-item">',
            '  <div class="fact-heading">',
            f'    {render_icon(item["icon"], css_class="fact-icon")}',
            f'    <p class="fact-label">{html_text(item["label"])}</p>',
            "  </div>",
            '  <p class="fact-value">',
        ]
        value_lines = [line.strip() for line in str(item["value"]).splitlines() if line.strip()]
        if len(value_lines) > 1:
            lines.extend(f'    <span class="fact-value-line">{html_text(line)}</span>' for line in value_lines)
        else:
            lines.append(f"    {html_text(item['value'])}")
        lines.extend(["  </p>", "</article>"])
        cards.append("\n".join(lines))
    return "\n          ".join(cards)


def render_stack(items: list[str], css_class: str = "stack-chip") -> str:
    return "".join(f'<span class="{css_class}">{html_text(item)}</span>' for item in items)


def render_experience(section_data: dict[str, Any], labels_data: dict[str, str], lang: str) -> str:
    articles = []
    for item in sorted(
        section_data["items"],
        key=lambda value: resume_date_upper_bound(value["startDate"]),
        reverse=True,
    ):
        company_icon = (
            render_icon(item["icon"], css_class="company-icon", alt=f"{item['company']} icon", hidden=False)
            if item.get("icon")
            else ""
        )
        if item.get("url"):
            company = (
                f'<a class="company-link" href="{html_attr(item["url"])}" target="_blank" '
                f'rel="noopener noreferrer">{html_text(item["company"])}</a>'
            )
            site_link = (
                f'<a class="company-site-link" href="{html_attr(item["url"])}" target="_blank" '
                f'rel="noopener noreferrer">{render_icon("external-link.svg")}<span>'
                f'{html_text(section_data["companySiteLabel"])}</span></a>'
            )
        else:
            company = f'<span class="company-link">{html_text(item["company"])}</span>'
            site_link = ""
        highlights = "\n".join(f"          <li>{html_text(bullet)}</li>" for bullet in item["highlights"])
        work_tags = "".join(
            f'<span class="work-tag">{html_text(value)}</span>'
            for value in experience_tag_values(item)
        )
        work_tags_block = f'<div class="work-tags">{work_tags}</div>' if work_tags else ""
        articles.append(
            f"""<article class="timeline-item">
        <div class="timeline-company-row">
          <div class="timeline-company-main">{company_icon}{company}{work_tags_block}</div>
          {site_link}
        </div>
        <h3 class="timeline-role">{html_text(item["role"])}</h3>
        <p class="timeline-meta">{html_text(format_period(item, labels_data, lang))}</p>
        <p class="timeline-intro">{html_text(item["summary"])}</p>
        <ul class="timeline-list">
{highlights}
        </ul>
        <div class="stack-list">{render_stack(item["stack"])}</div>
      </article>"""
        )
    return "\n      ".join(articles)


def render_education(section_data: dict[str, Any], labels_data: dict[str, str], lang: str) -> str:
    cards = []
    for item in section_data["items"]:
        icon = (
            render_icon(item["icon"], css_class="company-icon", alt=f"{item['institution']} icon", hidden=False)
            if item.get("icon")
            else ""
        )
        if item.get("url"):
            institution = (
                f'<a class="company-link" href="{html_attr(item["url"])}" target="_blank" '
                f'rel="noopener noreferrer">{html_text(item["institution"])}</a>'
            )
            site_link = (
                f'<a class="company-site-link" href="{html_attr(item["url"])}" target="_blank" '
                f'rel="noopener noreferrer">{render_icon("external-link.svg")}<span>'
                f'{html_text(section_data.get("institutionSiteLabel", "University site"))}</span></a>'
            )
        else:
            institution = f'<span class="company-link">{html_text(item["institution"])}</span>'
            site_link = ""
        cards.append(
            f"""<article class="education-item">
        <div class="timeline-company-row">
          <h3 class="education-heading timeline-company-main">{icon}{institution}</h3>
          {site_link}
        </div>
        <p class="education-degree">{html_text(item["degree"])}</p>
        <p class="education-meta">{html_text(format_education_period(item, labels_data, lang))}</p>
      </article>"""
        )
    return "\n      ".join(cards)


def render_strengths(cards: list[dict[str, str]]) -> str:
    return "\n          ".join(
        f'<article class="strength-card"><h3>{html_text(card["title"])}</h3><p>{html_text(card["body"])}</p></article>'
        for card in cards
    )


def render_skills(groups: list[dict[str, Any]]) -> str:
    cards = []
    for group in groups:
        items = "".join(f'<span class="skill-item">{html_text(item)}</span>' for item in group["items"])
        cards.append(f'<article class="skill-card"><h3>{html_text(group["title"])}</h3><div class="skill-items">{items}</div></article>')
    return "\n          ".join(cards)


def render_preferences(items: list[str]) -> str:
    return "\n          ".join(f"<li>{html_text(item)}</li>" for item in items)


def render_resume_actions(source: dict[str, Any], lang: str) -> str:
    file_names = download_file_names(source)
    labels_data = localized_tree(source["siteUi"]["resumeDownloads"]["labels"], lang, source["languages"])
    icons = {"pdf": "pdf.svg", "docx": "docx.svg", "txt": "txt.svg"}
    cards = []
    for extension in ("pdf", "docx", "txt"):
        file_name = file_names[extension]
        cards.append(
            f"""<a id="resume-{extension}" class="resume-button" href="{html_attr(download_path(source, lang, file_name))}" download="{html_attr(file_name)}" data-analytics-goal="resume_download_click" data-analytics-label="{extension.upper()} resume" data-analytics-section="resume">
          <img src="{html_attr(public_media_path(icons[extension]))}" alt="{extension.upper()} icon" />
          <div>
            <span class="format">{extension.upper()}</span>
            <span id="resume-{extension}-label" class="label">{html_text(labels_data[extension])}</span>
          </div>
        </a>"""
        )
    return "\n        ".join(cards)


def json_ld(source: dict[str, Any], data: dict[str, Any], lang: str, canonical_url: str) -> str:
    contacts = data["contacts"]
    same_as = [
        contact_value(contacts[key])
        for key in ("linkedin", "github", "telegram", "website")
        if key in contacts and is_web_url(contact_value(contacts[key]))
    ]
    person_id = f"{canonical_url}#person"
    person_node = {
        "@type": "Person",
        "@id": person_id,
        "name": data["person"]["name"],
        "jobTitle": data["person"]["headline"],
        "description": data["person"]["role"],
        "url": contact_value(contacts["website"]),
        "email": contact_value(contacts["email"]),
        "image": absolute_site_url(source, public_media_path(data["person"]["photo"]["src"])),
        "sameAs": same_as,
    }
    profile_node = {
        "@type": "ProfilePage",
        "@id": f"{canonical_url}#profile",
        "url": canonical_url,
        "name": f"{data['person']['name']} | {data['person']['headline']}",
        "inLanguage": lang,
        "about": {"@id": person_id},
        "mainEntity": {"@id": person_id},
    }
    return json.dumps({"@context": "https://schema.org", "@graph": [person_node, profile_node]}, ensure_ascii=False)


def script_json(value: str) -> str:
    return value.replace("<", "\\u003C").replace("</", "<\\/")


def safe_html(value: str) -> Markup:
    return Markup(value)


@lru_cache(maxsize=1)
def template_environment() -> Environment:
    return Environment(
        loader=FileSystemLoader(TEMPLATE_ROOT),
        autoescape=select_autoescape(("html", "j2")),
        keep_trailing_newline=True,
    )


def render_template(template_name: str, **context: Any) -> str:
    return template_environment().get_template(template_name).render(**context)


def alternate_links_html(source: dict[str, Any]) -> str:
    links = [
        f'    <link rel="alternate" hreflang="{item_lang}" href="{html_attr(page_url(source, item_lang))}" />'
        for item_lang in source["languages"]
    ]
    links.append(
        f'    <link rel="alternate" hreflang="x-default" href="{html_attr(page_url(source, source["defaultLanguage"]))}" />'
    )
    return "\n".join(links)


def generate_root_resolver_html(source: dict[str, Any]) -> str:
    default_lang = source["defaultLanguage"]
    data = localized_tree(source, default_lang, source["languages"])
    title = f"{data['person']['name']} - {data['siteUi']['navResume']}"
    page_paths = {lang: relative_public_path(page_path(source, lang), depth=0) for lang in source["languages"]}
    language_links = [{"code": lang.upper(), "path": page_paths[lang]} for lang in source["languages"]]
    return render_template(
        "root_resolver.html.j2",
        default_lang=default_lang,
        canonical_url=page_url(source, default_lang),
        alternate_links=safe_html(alternate_links_html(source)),
        title=title,
        pages_json=safe_html(script_json(json.dumps(page_paths, ensure_ascii=False, sort_keys=True))),
        default_lang_json=safe_html(script_json(json.dumps(default_lang))),
        default_path=page_paths[default_lang],
        language_links=language_links,
    )


def generate_site_html(source: dict[str, Any], lang: str) -> str:
    data = localized_tree(source, lang, source["languages"])
    site_ui = source["siteUi"]
    profile = data["person"]
    labels_data = data["resumeLabels"]
    title = f"{profile['name']} | {profile['headline']}"
    canonical_url = page_url(source, lang)
    cover_url = absolute_site_url(source, SITE_COVER_IMAGE)
    theme_labels = localized_tree(site_ui["theme"], lang, source["languages"])
    description = profile["role"]
    og_locale = "ru_RU" if lang == "ru" else "en_US"
    hero_style = ""
    if profile["photo"].get("position"):
        hero_style += f"object-position: {profile['photo']['position']};"
    if profile["photo"].get("filter"):
        hero_style += f"--photo-filter: {profile['photo']['filter']};"
    hero_sizes = "(max-width: 980px) min(100vw, 340px), 360px"
    footer_text = localized_tree(site_ui["footer"], lang, source["languages"]).replace(
        "{year}", str(date.today().year)
    ).replace("{name}", profile["name"])
    contact_links = "\n              ".join(
        render_contact_link(key, data["contacts"][key])
        for key in ("email", "linkedin", "github", "telegram")
    )
    return render_template(
        "resume_page.html.j2",
        lang=lang,
        canonical_path=page_path(source, lang),
        description=description,
        canonical_url=canonical_url,
        alternate_links=safe_html(alternate_links_html(source)),
        cover_url=cover_url,
        hero_preload_href=responsive_variant_path(
            profile["photo"]["src"],
            HERO_PRELOAD_WIDTH,
            HERO_PRELOAD_FORMAT,
        ),
        hero_preload_srcset=responsive_srcset(
            profile["photo"]["src"],
            HERO_PRELOAD_FORMAT,
        ),
        hero_sizes=hero_sizes,
        og_locale=og_locale,
        site_host=site_host(source),
        title=title,
        json_ld_json=safe_html(script_json(json_ld(source, data, lang, canonical_url))),
        embedded_resume_data_json=safe_html(
            script_json(
                json.dumps(runtime_page_data(source, lang), ensure_ascii=False, sort_keys=True, separators=(",", ":"))
            )
        ),
        yandex_metrika_id=YANDEX_METRIKA_ID,
        yandex_metrika_origin=YANDEX_METRIKA_ORIGIN,
        lang_switcher_label=localized_tree(site_ui["langSwitcherLabel"], lang, source["languages"]),
        language_switch=safe_html(render_language_switch(source, lang)),
        nav_resume=localized_tree(site_ui["navResume"], lang, source["languages"]),
        experience_title=data["experience"]["title"],
        education_title=data["education"]["title"],
        strengths_title=data["strengths"]["title"],
        stack_label=labels_data["stack"],
        theme_switcher_label=theme_labels["switcherLabel"],
        theme_to_light=theme_labels["toLight"],
        theme_to_dark=theme_labels["toDark"],
        sun_icon=safe_html(render_icon("sun.svg")),
        moon_icon=safe_html(render_icon("moon.svg")),
        headline=profile["headline"],
        person_name=profile["name"],
        role=profile["role"],
        summary=profile["summary"],
        contact_links=safe_html(contact_links),
        hero_photo=safe_html(
            responsive_picture_html(
                profile["photo"]["src"],
                css_class="hero-photo",
                alt=profile["photo"]["caption"],
                sizes=hero_sizes,
                loading="eager",
                fetchpriority="high",
                image_id="hero-photo",
                style=hero_style or None,
            )
        ),
        hero_photo_caption=profile["photo"]["caption"],
        facts=safe_html(render_facts(profile["facts"])),
        download_panel_icon=safe_html(render_panel_icon("download.svg")),
        resume_downloads_title=localized_tree(site_ui["resumeDownloads"]["title"], lang, source["languages"]),
        resume_actions=safe_html(render_resume_actions(source, lang)),
        briefcase_panel_icon=safe_html(render_panel_icon("briefcase.svg")),
        experience_items=safe_html(render_experience(data["experience"], labels_data, lang)),
        education_panel_icon=safe_html(render_panel_icon("education.svg")),
        education_subtitle=data["education"]["subtitle"],
        education_items=safe_html(render_education(data["education"], labels_data, lang)),
        star_panel_icon=safe_html(render_panel_icon("star.svg")),
        strengths_subtitle=data["strengths"]["subtitle"],
        strengths_cards=safe_html(render_strengths(data["strengths"]["cards"])),
        layers_panel_icon=safe_html(render_panel_icon("layers.svg")),
        skills_title=data["skills"]["title"],
        skills_groups=safe_html(render_skills(data["skills"]["groups"])),
        preferences_title=data["preferences"]["title"],
        preferences_items=safe_html(render_preferences(data["preferences"]["items"])),
        footer_text=footer_text,
    )


def write_static_site(source: dict[str, Any], output_root: Path) -> None:
    output_root.mkdir(parents=True, exist_ok=True)
    (output_root / "index.html").write_text(generate_root_resolver_html(source), encoding="utf-8")
    for lang in source["languages"]:
        lang_dir = output_root / lang
        lang_dir.mkdir(parents=True, exist_ok=True)
        (lang_dir / "index.html").write_text(generate_site_html(source, lang), encoding="utf-8")


def generate_outputs(
    source: dict[str, Any],
    output_root: Path,
    build_dir: Path,
) -> None:
    generate_responsive_media(source)
    write_static_site(source, output_root)
    file_names = download_file_names(source)
    for lang in source["languages"]:
        lang_dir = output_root / lang
        generate_txt(source, lang, lang_dir / file_names["txt"])
        download_scale = generate_pdf(source, lang, lang_dir / file_names["pdf"], build_dir)
        docx_scale = download_scale * DOWNLOAD_STYLES["layout"]["docxScaleAdjustment"].get(lang, 1.0)
        generate_docx(source, lang, lang_dir / file_names["docx"], build_dir, docx_scale)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--data", type=Path, default=ROOT / "resume/resume.yaml")
    parser.add_argument("--output-root", type=Path, default=ROOT)
    parser.add_argument("--build-dir", type=Path, default=ROOT / ".build/resume-assets")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data = load_data(args.data)
    generate_outputs(data, args.output_root, args.build_dir)
    print("Resume outputs generated successfully.")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # pragma: no cover - keeps CI/local failures readable.
        print(f"Resume generation failed: {exc}", file=sys.stderr)
        raise
